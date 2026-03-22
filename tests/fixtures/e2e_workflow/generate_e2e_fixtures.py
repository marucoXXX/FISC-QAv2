"""E2Eワークフローテスト用フィクスチャ生成スクリプト

XLSX / DOCX 両形式の質問票と、過去回答・共通回答のJSONを生成する。

  python tests/fixtures/e2e_workflow/generate_e2e_fixtures.py
"""
from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook

FIXTURES_DIR = Path(__file__).resolve().parent

# ---------------------------------------------------------------------------
# 質問データ（10問）
# Q1-Q5: 過去回答マッチ対象
# Q6-Q8: 共通回答マッチ対象
# Q9-Q10: KB生成対象（新規質問）
# ---------------------------------------------------------------------------

QUESTIONS = [
    # --- 過去回答マッチ対象 (Q1-Q5) ---
    {"no": 1, "major": "セキュリティ管理", "minor": "ポリシー管理",
     "question": "情報セキュリティポリシーの策定・承認状況を教えてください"},
    {"no": 2, "major": "アクセス管理", "minor": "ユーザー認証",
     "question": "多要素認証（MFA）の導入状況について説明してください"},
    {"no": 3, "major": "ネットワーク管理", "minor": "通信暗号化",
     "question": "通信経路のTLS暗号化の適用状況を説明してください"},
    {"no": 4, "major": "バックアップ管理", "minor": "バックアップ",
     "question": "バックアップの取得方式と遠隔地保管の状況を説明してください"},
    {"no": 5, "major": "インシデント対応", "minor": "対応手順",
     "question": "セキュリティインシデント発生時の対応手順を説明してください"},
    # --- 共通回答マッチ対象 (Q6-Q8) ---
    {"no": 6, "major": "物理セキュリティ", "minor": "入退室管理",
     "question": "サーバルームの入退室はどのように管理されていますか"},
    {"no": 7, "major": "教育・訓練", "minor": "セキュリティ教育",
     "question": "従業員のセキュリティ教育はどのように実施していますか"},
    {"no": 8, "major": "変更管理", "minor": "変更手順",
     "question": "本番環境への変更管理プロセスを教えてください"},
    # --- KB生成対象 (Q9-Q10) ---
    {"no": 9, "major": "システム開発", "minor": "コード品質",
     "question": "ソースコードの品質管理とセキュリティレビューの体制を説明してください"},
    {"no": 10, "major": "ログ・監視", "minor": "リアルタイム監視",
     "question": "SIEMによるリアルタイム監視体制について説明してください"},
]

# ---------------------------------------------------------------------------
# 過去回答データ（5件）— Q1-Q5と高類似度の質問文 + 回答
# 質問文は微調整してあるがコサイン類似度 >= 0.7 を保証
# ---------------------------------------------------------------------------

PAST_ANSWERS = [
    {
        "question_text": "情報セキュリティポリシーの策定と承認の状況を説明してください",
        "answer_text": "情報セキュリティポリシー v3.2に基づき運用中。年次レビューを2024年4月に完了。",
    },
    {
        "question_text": "多要素認証（MFA）の導入と運用状況について説明してください",
        "answer_text": "全重要システムにMFA導入済み（TOTP方式）。VPN接続時も必須化。",
    },
    {
        "question_text": "通信経路におけるTLS暗号化の適用状況について説明してください",
        "answer_text": "全外部通信および社内重要通信にTLS 1.3を適用済み。TLS 1.0/1.1は全面廃止完了。",
    },
    {
        "question_text": "バックアップの取得方式および遠隔地保管の状況について説明してください",
        "answer_text": "3-2-1ルールに基づくバックアップ体制を運用中。四半期ごとのリストアテスト実施。RTO=4時間、RPO=1時間。",
    },
    {
        "question_text": "セキュリティインシデント発生時の初動対応手順について説明してください",
        "answer_text": "CSIRT体制を構築済み。インシデント対応手順書(IRP-2024-v2)に基づき、年4回の訓練を実施。",
    },
]

# ---------------------------------------------------------------------------
# 共通回答データ（3件）— Q6-Q8に意味的マッチ
# ---------------------------------------------------------------------------

COMMON_ANSWERS = [
    {
        "question_pattern": "データセンター・サーバルームへの入退室管理の方法",
        "answer_text": "データセンターはISO 27001認証取得済み施設を利用。入退室はICカード＋生体認証の二要素で管理。監視カメラ映像のAI異常検知を運用中。",
        "category": "物理セキュリティ",
    },
    {
        "question_pattern": "全従業員を対象としたセキュリティ教育の実施方法",
        "answer_text": "全従業員向けeラーニング（年2回必須）および役職別専門研修を実施。受講率99.2%。標的型メール訓練を四半期ごとに実施。",
        "category": "教育・訓練",
    },
    {
        "question_pattern": "本番環境に対する変更管理プロセスの運用手順",
        "answer_text": "変更管理プロセス(CHG-PROC-v3)に基づき、CAB（変更諮問委員会）による承認制を運用。ServiceNowで一元管理。",
        "category": "変更管理",
    },
]


# =========================================================================
# 生成関数
# =========================================================================

def _make_xlsx(path: Path):
    """質問票をXLSX形式で生成"""
    wb = Workbook()
    ws = wb.active
    ws.title = "FISC回答"
    ws.append(["No.", "大分類", "小分類", "質問内容", "回答"])
    for q in QUESTIONS:
        ws.append([q["no"], q["major"], q["minor"], q["question"], ""])
    wb.save(str(path))


def _make_docx(path: Path):
    """質問票をDOCX（テーブル形式）で生成"""
    doc = Document()
    doc.add_heading("FISC質問票", level=0)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["No.", "大分類", "小分類", "質問内容", "回答"]):
        hdr[i].text = h
    for q in QUESTIONS:
        row = table.add_row().cells
        row[0].text = str(q["no"])
        row[1].text = q["major"]
        row[2].text = q["minor"]
        row[3].text = q["question"]
        row[4].text = ""
    doc.save(str(path))


def generate_all():
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)

    print("=== E2Eワークフロー テストフィクスチャ生成 ===\n")

    # XLSX質問票
    xlsx_path = FIXTURES_DIR / "questionnaire_10q.xlsx"
    _make_xlsx(xlsx_path)
    print(f"  -> {xlsx_path.name}")

    # DOCX質問票
    docx_path = FIXTURES_DIR / "questionnaire_10q.docx"
    _make_docx(docx_path)
    print(f"  -> {docx_path.name}")

    # 過去回答JSON
    pa_path = FIXTURES_DIR / "past_answers.json"
    pa_path.write_text(json.dumps(PAST_ANSWERS, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  -> {pa_path.name}")

    # 共通回答JSON
    ca_path = FIXTURES_DIR / "common_answers.json"
    ca_path.write_text(json.dumps(COMMON_ANSWERS, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  -> {ca_path.name}")

    print("\n=== 生成完了 (4 files) ===")


if __name__ == "__main__":
    generate_all()
