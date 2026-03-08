"""FISC-QAv2 テストフィクスチャ生成スクリプト

ダミーの PDF / DOCX / Excel ファイルを一括生成する。
コンテンツ定義は content/ サブモジュールに分離されている。

  python tests/fixtures/generate_fixtures.py
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# ---------------------------------------------------------------------------
# コンテンツインポート
# ---------------------------------------------------------------------------
from content.questions import QUESTIONS, NUM_QUESTIONS, _MAJOR_TO_SOURCES
from content.policies import POLICY_FILES
from content.system_docs import SYSTEM_DOC_FILES
from content.operations import OPERATIONS_FILES, CHANGE_LOG_HEADERS, CHANGE_LOG_ROWS
from content.regulations import REGULATION_FILES
from content.past_answers import PA_HEADERS, generate_past_answers_2024, generate_past_answers_2023

# ---------------------------------------------------------------------------
# パス定義
# ---------------------------------------------------------------------------
FIXTURES = Path(__file__).resolve().parent
INPUT_DIR = FIXTURES / "input"
KB_DIR = FIXTURES / "kb"
EXPECTED_DIR = FIXTURES / "expected"

POLICIES_DIR = KB_DIR / "policies"
PAST_ANSWERS_DIR = KB_DIR / "past_answers"
SYSTEM_DOCS_DIR = KB_DIR / "system_docs"
OPERATIONS_DIR = KB_DIR / "operations"
REGULATIONS_DIR = KB_DIR / "regulations"

ALL_DIRS = [
    INPUT_DIR,
    POLICIES_DIR,
    PAST_ANSWERS_DIR,
    SYSTEM_DOCS_DIR,
    OPERATIONS_DIR,
    REGULATIONS_DIR,
    EXPECTED_DIR,
]

# ---------------------------------------------------------------------------
# 日本語フォント登録（macOS 標準ヒラギノ / フォールバック）
# ---------------------------------------------------------------------------
_JP_FONT_NAME = "IPAGothic"
_JP_FONT_REGISTERED = False

_CANDIDATE_FONTS = [
    # macOS — TrueType outlines (CFF/PostScript outlines in TTC are unsupported by ReportLab)
    "/Library/Fonts/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/NotoSansGothic-Regular.ttf",
    "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
    # Linux
    "/usr/share/fonts/truetype/ipafont-gothic/ipag.ttf",
    "/usr/share/fonts/opentype/ipafont-gothic/ipag.ttf",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
]


def _register_jp_font() -> str:
    """日本語フォントを登録して使用可能なフォント名を返す。"""
    global _JP_FONT_REGISTERED, _JP_FONT_NAME
    if _JP_FONT_REGISTERED:
        return _JP_FONT_NAME
    for fp in _CANDIDATE_FONTS:
        p = Path(fp)
        if p.exists():
            try:
                if fp.endswith(".ttc"):
                    pdfmetrics.registerFont(TTFont(_JP_FONT_NAME, str(p), subfontIndex=0))
                else:
                    pdfmetrics.registerFont(TTFont(_JP_FONT_NAME, str(p)))
                _JP_FONT_REGISTERED = True
                return _JP_FONT_NAME
            except Exception:
                continue
    # フォールバック: Helvetica（日本語は化けるが生成は成功する）
    _JP_FONT_NAME = "Helvetica"
    _JP_FONT_REGISTERED = True
    return _JP_FONT_NAME


# =========================================================================
# PDF 生成
# =========================================================================

def _make_pdf(path: Path, title: str, sections: list[tuple[str, str]], target_pages: int = 6):
    """ReportLab で日本語テキスト入りの PDF を生成する。"""
    font_name = _register_jp_font()
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            topMargin=20 * mm, bottomMargin=20 * mm)
    styles = {
        "title": ParagraphStyle("title", fontName=font_name, fontSize=16, leading=22,
                                spaceAfter=12),
        "heading": ParagraphStyle("heading", fontName=font_name, fontSize=13, leading=18,
                                  spaceAfter=8, spaceBefore=12),
        "body": ParagraphStyle("body", fontName=font_name, fontSize=10, leading=15,
                               spaceAfter=6),
    }

    story: list = []
    story.append(Paragraph(title, styles["title"]))
    story.append(Spacer(1, 6 * mm))

    # セクションを繰り返してページ数を稼ぐ
    repeat = max(1, target_pages // max(len(sections), 1))
    for _ in range(repeat):
        for heading, body in sections:
            story.append(Paragraph(heading, styles["heading"]))
            for para in body.split("\n\n"):
                story.append(Paragraph(para.replace("\n", "<br/>"), styles["body"]))
            story.append(Spacer(1, 3 * mm))

    doc.build(story)


# =========================================================================
# DOCX 生成
# =========================================================================

def _make_docx(path: Path, title: str, sections: list[tuple[str, str]]):
    doc = Document()
    doc.add_heading(title, level=0)
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        for para in body.split("\n\n"):
            doc.add_paragraph(para)
    doc.save(str(path))


# =========================================================================
# Excel 生成
# =========================================================================

def _make_excel(path: Path, headers: list[str], rows: list[list], sheet_name: str = "Sheet1"):
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(str(path))


# =========================================================================
# ファイル生成ヘルパー
# =========================================================================

def _generate_file_set(base_dir: Path, file_map: dict, label: str):
    """PDF/DOCX ファイルセットを生成する汎用ヘルパー。"""
    for fname, meta in file_map.items():
        path = base_dir / fname
        if fname.endswith(".pdf"):
            _make_pdf(path, meta["title"], meta["sections"],
                      target_pages=meta.get("pages", 6))
        else:
            _make_docx(path, meta["title"], meta["sections"])
        print(f"  -> {path.relative_to(FIXTURES)}")


# =========================================================================
# メイン生成関数
# =========================================================================

def generate_all():
    print("=== FISC-QAv2 テストフィクスチャ生成 ===\n")

    # ディレクトリ作成
    for d in ALL_DIRS:
        d.mkdir(parents=True, exist_ok=True)

    # --- 1. 入力ファイル: questionnaire.xlsx (回答欄は空) ---
    print("[1/8] questionnaire.xlsx ...")
    headers = ["No.", "大分類", "小分類", "質問内容", "回答", "対応状況", "根拠ソース", "確信度", "備考"]
    rows = [[q["no"], q["major"], q["minor"], q["question"], "", "", "", "", ""]
            for q in QUESTIONS]
    _make_excel(INPUT_DIR / "questionnaire.xlsx", headers, rows, sheet_name="FISC回答結果")

    # --- 2. KB: policies (5 files) ---
    print("[2/8] KB policies ...")
    _generate_file_set(POLICIES_DIR, POLICY_FILES, "policies")

    # --- 3. KB: system_docs (5 files) ---
    print("[3/8] KB system_docs ...")
    _generate_file_set(SYSTEM_DOCS_DIR, SYSTEM_DOC_FILES, "system_docs")

    # --- 4. KB: operations (3 files: 2 PDF/DOCX + 1 Excel) ---
    print("[4/8] KB operations ...")
    _generate_file_set(OPERATIONS_DIR, OPERATIONS_FILES, "operations")
    # change_management_log.xlsx は別途 Excel として生成
    change_log_path = OPERATIONS_DIR / "change_management_log.xlsx"
    _make_excel(change_log_path, CHANGE_LOG_HEADERS, CHANGE_LOG_ROWS, sheet_name="変更管理台帳")
    print(f"  -> {change_log_path.relative_to(FIXTURES)}")

    # --- 5. KB: regulations (2 files) ---
    print("[5/8] KB regulations ...")
    _generate_file_set(REGULATIONS_DIR, REGULATION_FILES, "regulations")

    # --- 6. KB: past_answers (2 files) ---
    print("[6/8] KB past_answers ...")
    _make_excel(PAST_ANSWERS_DIR / "past_answer_2024.xlsx", PA_HEADERS,
                generate_past_answers_2024(), sheet_name="回答実績")
    _make_excel(PAST_ANSWERS_DIR / "past_answer_2023.xlsx", PA_HEADERS,
                generate_past_answers_2023(), sheet_name="回答実績")
    print("  -> past_answer_2024.xlsx, past_answer_2023.xlsx")

    # --- 7. 期待出力 JSON ---
    print("[7/8] Expected outputs ...")
    _generate_expected_outputs()

    # --- 8. 完了サマリ ---
    print("\n=== 生成完了 ===")
    print(f"合計ファイル数: {_count_files()}")


def _generate_expected_outputs():
    # index.json (17 entries)
    kb_files = []
    for subdir in ["policies", "past_answers", "system_docs", "operations", "regulations"]:
        d = KB_DIR / subdir
        if d.exists():
            for f in sorted(d.iterdir()):
                if f.is_file():
                    kb_files.append({
                        "file_name": f.name,
                        "path": f"{subdir}/{f.name}",
                        "category": subdir,
                        "summary": f"[要約] {f.stem} の内容",
                        "estimated_tokens": 2000,
                    })
    index_path = EXPECTED_DIR / "index.json"
    index_path.write_text(json.dumps(kb_files, ensure_ascii=False, indent=2), encoding="utf-8")

    # routing_map.json (30 entries)
    routing = {}
    for q in QUESTIONS:
        no = q["no"]
        major = q["major"]
        sources = _MAJOR_TO_SOURCES.get(major, [])
        routing[str(no)] = {
            "major": major,
            "minor": q["minor"],
            "assigned_sources": sources,
        }
    routing_path = EXPECTED_DIR / "routing_map.json"
    routing_path.write_text(json.dumps(routing, ensure_ascii=False, indent=2), encoding="utf-8")

    # final_answers.json (30 entries)
    answers = {}
    for q in QUESTIONS:
        no = q["no"]
        answers[str(no)] = {
            "question": q["question"],
            "answer": "",
            "status": "",
            "evidence": "",
            "confidence": "",
            "needs_review": True,
        }
    answers_path = EXPECTED_DIR / "final_answers.json"
    answers_path.write_text(json.dumps(answers, ensure_ascii=False, indent=2), encoding="utf-8")


def _count_files() -> int:
    count = 0
    for d in [INPUT_DIR, POLICIES_DIR, PAST_ANSWERS_DIR, SYSTEM_DOCS_DIR,
              OPERATIONS_DIR, REGULATIONS_DIR, EXPECTED_DIR]:
        if d.exists():
            count += sum(1 for f in d.iterdir() if f.is_file())
    return count


if __name__ == "__main__":
    generate_all()
