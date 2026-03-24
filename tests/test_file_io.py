"""Tests for the file I/O module (read/write questionnaires)."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest

from src.file_io import FormatConfig, Question, _col_letter_to_index, read_questionnaire, write_answers_to_original


def _make_xlsx_questionnaire(tmp_path: Path, q_col: str = "D", rows: list[list] | None = None) -> Path:
    """Create a test questionnaire Excel file."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["No.", "大分類", "小分類", "質問内容", "回答"])
    if rows is None:
        rows = [
            [1, "セキュリティ", "認証", "MFAの導入状況は？", ""],
            [2, "バックアップ", "RPO", "RPOの設定は？", ""],
            [3, "アクセス管理", "権限", "最小権限の原則は？", ""],
        ]
    for r in rows:
        ws.append(r)

    path = tmp_path / "questionnaire.xlsx"
    wb.save(str(path))
    wb.close()
    return path


def _make_docx_questionnaire(tmp_path: Path) -> Path:
    """Create a test questionnaire Word file with a table."""
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=4, cols=5)
    # Header row
    table.rows[0].cells[0].text = "No."
    table.rows[0].cells[1].text = "大分類"
    table.rows[0].cells[2].text = "小分類"
    table.rows[0].cells[3].text = "質問内容"
    table.rows[0].cells[4].text = "回答"
    # Data rows
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "セキュリティ"
    table.rows[1].cells[2].text = "認証"
    table.rows[1].cells[3].text = "MFAの導入状況は？"
    table.rows[1].cells[4].text = ""

    table.rows[2].cells[0].text = "2"
    table.rows[2].cells[1].text = "バックアップ"
    table.rows[2].cells[2].text = "RPO"
    table.rows[2].cells[3].text = "RPOの設定は？"
    table.rows[2].cells[4].text = ""

    table.rows[3].cells[0].text = "3"
    table.rows[3].cells[1].text = "アクセス管理"
    table.rows[3].cells[2].text = "権限"
    table.rows[3].cells[3].text = "最小権限の原則は？"
    table.rows[3].cells[4].text = ""

    path = tmp_path / "questionnaire.docx"
    doc.save(str(path))
    return path


class TestColLetterToIndex:
    def test_col_letter_to_index(self):
        assert _col_letter_to_index("A") == 0
        assert _col_letter_to_index("B") == 1
        assert _col_letter_to_index("D") == 3
        assert _col_letter_to_index("Z") == 25
        assert _col_letter_to_index("AA") == 26


class TestReadXlsx:
    def test_read_xlsx_questionnaire(self, tmp_path):
        path = _make_xlsx_questionnaire(tmp_path)
        fc = FormatConfig(file_format="xlsx", question_col="D", answer_col="E",
                          header_row=1, data_start_row=2)
        questions = read_questionnaire(path, fc)
        assert len(questions) == 3
        assert questions[0].question_text == "MFAの導入状況は？"
        assert questions[0].question_no == 1
        assert questions[1].major == "バックアップ"

    def test_read_xlsx_custom_columns(self, tmp_path):
        """Test reading with non-standard column layout (question in B, answer in F)."""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["No.", "質問", "C", "D", "E", "回答"])
        ws.append([1, "MFAは？", "", "", "", ""])
        ws.append([2, "RPOは？", "", "", "", ""])

        path = tmp_path / "custom.xlsx"
        wb.save(str(path))
        wb.close()

        fc = FormatConfig(file_format="xlsx", question_col="B", answer_col="F",
                          header_row=1, data_start_row=2)
        questions = read_questionnaire(path, fc)
        assert len(questions) == 2
        assert questions[0].question_text == "MFAは？"


class TestWriteXlsx:
    def test_write_xlsx_answers(self, tmp_path):
        path = _make_xlsx_questionnaire(tmp_path)
        fc = FormatConfig(file_format="xlsx", question_col="D", answer_col="E",
                          header_row=1, data_start_row=2)
        answers = {1: "導入済み", 3: "適用済み"}
        output = tmp_path / "output.xlsx"

        write_answers_to_original(path, answers, fc, output)

        from openpyxl import load_workbook
        wb = load_workbook(str(output))
        ws = wb.active
        assert ws["E2"].value == "導入済み"
        assert ws["E3"].value is None  # Q2 has no answer
        assert ws["E4"].value == "適用済み"
        wb.close()


class TestReadDocx:
    def test_read_docx_questionnaire(self, tmp_path):
        path = _make_docx_questionnaire(tmp_path)
        fc = FormatConfig(file_format="docx", question_col="D", answer_col="E",
                          header_row=1, data_start_row=2, table_index=0)
        questions = read_questionnaire(path, fc)
        assert len(questions) == 3
        assert questions[0].question_text == "MFAの導入状況は？"

    def test_read_docx_invalid_table_index(self, tmp_path):
        path = _make_docx_questionnaire(tmp_path)
        fc = FormatConfig(file_format="docx", question_col="D", answer_col="E",
                          header_row=1, data_start_row=2, table_index=99)
        questions = read_questionnaire(path, fc)
        assert questions == []


class TestWriteDocx:
    def test_write_docx_answers(self, tmp_path):
        path = _make_docx_questionnaire(tmp_path)
        fc = FormatConfig(file_format="docx", question_col="D", answer_col="E",
                          header_row=1, data_start_row=2, table_index=0)
        answers = {1: "導入済み", 2: "4時間"}
        output = tmp_path / "output.docx"

        write_answers_to_original(path, answers, fc, output)

        from docx import Document
        doc = Document(str(output))
        table = doc.tables[0]
        assert table.rows[1].cells[4].text == "導入済み"
        assert table.rows[2].cells[4].text == "4時間"


def _make_xlsx_multi_col(tmp_path: Path) -> Path:
    """Create xlsx with columns A-F: No, Cat, Sub, Question, Answer1, Answer2."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["No.", "分類", "小分類", "質問", "対応状況", "代替策"])
    ws.append([1, "セキュリティ", "認証", "MFA導入は？", "", ""])
    ws.append([2, "バックアップ", "RPO", "RPOは？", "", ""])
    path = tmp_path / "multi_col.xlsx"
    wb.save(str(path))
    wb.close()
    return path


class TestWriteXlsxNoFallback:
    """Fix 1: column_definitions 使用時に answer_text へのフォールバックが起きないことを検証."""

    def test_write_xlsx_no_fallback_with_col_defs(self, tmp_path):
        """per_col_answers のテキストのみが出力され、answers dict は使われない."""
        path = _make_xlsx_multi_col(tmp_path)
        fc = FormatConfig(file_format="xlsx", question_col="D", answer_col="E",
                          header_row=1, data_start_row=2)
        col_defs = [
            {"col": "D", "role": "question"},
            {"col": "E", "role": "answer"},
            {"col": "F", "role": "answer"},
        ]
        answers = {1: "fallback text", 2: "fallback text"}
        per_col = {1: {"E": "correct E", "F": "correct F"}, 2: {"E": "correct E2"}}
        output = tmp_path / "out.xlsx"

        write_answers_to_original(path, answers, fc, output,
                                  column_definitions=col_defs, per_col_answers=per_col)

        from openpyxl import load_workbook
        wb = load_workbook(str(output))
        ws = wb.active
        assert ws["E2"].value == "correct E"
        assert ws["F2"].value == "correct F"
        assert ws["E3"].value == "correct E2"
        assert ws["F3"].value is None  # not "fallback text"
        wb.close()

    def test_write_xlsx_empty_per_col_no_fallback(self, tmp_path):
        """per_col_answers={} の場合、answer_text にフォールバックしない."""
        path = _make_xlsx_multi_col(tmp_path)
        fc = FormatConfig(file_format="xlsx", question_col="D", answer_col="E",
                          header_row=1, data_start_row=2)
        col_defs = [
            {"col": "D", "role": "question"},
            {"col": "E", "role": "answer"},
            {"col": "F", "role": "answer"},
        ]
        answers = {1: "AI proposal text", 2: "AI proposal text"}
        per_col = {1: {}, 2: {}}  # 空 dict = 何も書かない
        output = tmp_path / "out.xlsx"

        write_answers_to_original(path, answers, fc, output,
                                  column_definitions=col_defs, per_col_answers=per_col)

        from openpyxl import load_workbook
        wb = load_workbook(str(output))
        ws = wb.active
        assert ws["E2"].value is None  # NOT "AI proposal text"
        assert ws["F2"].value is None
        assert ws["E3"].value is None
        assert ws["F3"].value is None
        wb.close()

    def test_write_xlsx_partial_col_texts(self, tmp_path):
        """一部列のみにテキストがある場合、他列は None."""
        path = _make_xlsx_multi_col(tmp_path)
        fc = FormatConfig(file_format="xlsx", question_col="D", answer_col="E",
                          header_row=1, data_start_row=2)
        col_defs = [
            {"col": "D", "role": "question"},
            {"col": "E", "role": "answer"},
            {"col": "F", "role": "answer"},
        ]
        answers = {1: "should not appear"}
        per_col = {1: {"E": "only E"}}  # F のキーなし
        output = tmp_path / "out.xlsx"

        write_answers_to_original(path, answers, fc, output,
                                  column_definitions=col_defs, per_col_answers=per_col)

        from openpyxl import load_workbook
        wb = load_workbook(str(output))
        ws = wb.active
        assert ws["E2"].value == "only E"
        assert ws["F2"].value is None  # F は per_col_answers にないので None
        wb.close()

    def test_write_xlsx_legacy_uses_answers(self, tmp_path):
        """レガシーモード(column_definitions なし)は従来通り answers dict を使用."""
        path = _make_xlsx_questionnaire(tmp_path)
        fc = FormatConfig(file_format="xlsx", question_col="D", answer_col="E",
                          header_row=1, data_start_row=2)
        answers = {1: "導入済み", 3: "適用済み"}
        output = tmp_path / "out.xlsx"

        # column_definitions=None → レガシーモード
        write_answers_to_original(path, answers, fc, output)

        from openpyxl import load_workbook
        wb = load_workbook(str(output))
        ws = wb.active
        assert ws["E2"].value == "導入済み"
        assert ws["E3"].value is None
        assert ws["E4"].value == "適用済み"
        wb.close()

    def test_write_docx_no_fallback_with_col_defs(self, tmp_path):
        """DOCX 版も column_definitions 使用時にフォールバックしない."""
        path = _make_docx_questionnaire(tmp_path)
        fc = FormatConfig(file_format="docx", question_col="D", answer_col="E",
                          header_row=1, data_start_row=2, table_index=0)
        col_defs = [
            {"col": "D", "role": "question"},
            {"col": "E", "role": "answer"},
        ]
        answers = {1: "fallback text", 2: "fallback text"}
        per_col = {1: {"E": "correct text"}, 2: {}}  # Q2 は空
        output = tmp_path / "out.docx"

        write_answers_to_original(path, answers, fc, output,
                                  column_definitions=col_defs, per_col_answers=per_col)

        from docx import Document
        doc = Document(str(output))
        table = doc.tables[0]
        assert table.rows[1].cells[4].text == "correct text"
        assert table.rows[2].cells[4].text == ""  # NOT "fallback text"
