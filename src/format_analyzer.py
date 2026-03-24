"""Format Analyzer: Excelプレビュー抽出 + LLMによるフォーマット自動解析."""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

import litellm
from litellm.utils import supports_response_schema

logger = logging.getLogger(__name__)

SYSTEM_PROMPT = """\
あなたは金融機関のアンケート（質問票）の構造を分析する専門家です。
提供されたExcelまたはWordファイルの先頭行データから、以下の列マッピングを特定してください。

分析対象:
1. question_col: 質問文が書かれている列のExcel列記号（例: "D"）
2. answer_col: 回答を記入すべき列のExcel列記号（例: "E"） - 通常空欄またはほぼ空欄
3. header_row: ヘッダー行の行番号（1始まり）
4. data_start_row: データが始まる行番号（1始まり）
5. format_type: "freetext"（自由記述）| "choices"（選択肢あり）| "assessment"（○/△/× 判定）
6. choices_col: 選択肢や判定記号の列（該当する場合のみ、なければ空文字）
7. remarks_col: 備考欄の列（該当する場合のみ、なければ空文字）

判断基準:
- 質問列: 長いテキストが入っている列で、質問文・確認事項のような内容
- 回答列: 質問列の隣またはその近くで、空欄が多い列（回答を記入する場所）
- ヘッダー行: 「質問」「回答」「項目」「確認事項」などの見出しがある行
- 選択肢列: 「はい/いいえ」「○/△/×」「該当/非該当」のような内容がある列
- 備考列: 「備考」「コメント」「補足」という見出しの列
- 大分類・小分類列: 質問列より左にある、カテゴリ分類のような短い文字列の列

以下のJSON形式で回答してください:
{
  "question_col": "D",
  "answer_col": "E",
  "header_row": 1,
  "data_start_row": 2,
  "format_type": "freetext",
  "choices_col": "",
  "remarks_col": "",
  "confidence": {
    "question_col": "high",
    "answer_col": "high",
    "header_row": "high",
    "data_start_row": "high",
    "format_type": "high",
    "choices_col": "low",
    "remarks_col": "low"
  },
  "reasoning": {
    "question_col": "列の内容に関する説明",
    "answer_col": "列の内容に関する説明",
    "header_row": "行の内容に関する説明",
    "data_start_row": "行の内容に関する説明",
    "format_type": "判断理由",
    "choices_col": "判断理由",
    "remarks_col": "判断理由"
  }
}

confidenceは "high" | "medium" | "low" のいずれかです。
"""

_RESPONSE_SCHEMA = {
    "type": "json_schema",
    "json_schema": {
        "name": "format_analysis",
        "strict": True,
        "schema": {
            "type": "object",
            "properties": {
                "question_col": {"type": "string"},
                "answer_col": {"type": "string"},
                "header_row": {"type": "integer"},
                "data_start_row": {"type": "integer"},
                "format_type": {"type": "string"},
                "choices_col": {"type": "string"},
                "remarks_col": {"type": "string"},
                "confidence": {
                    "type": "object",
                    "properties": {
                        "question_col": {"type": "string"},
                        "answer_col": {"type": "string"},
                        "header_row": {"type": "string"},
                        "data_start_row": {"type": "string"},
                        "format_type": {"type": "string"},
                        "choices_col": {"type": "string"},
                        "remarks_col": {"type": "string"},
                    },
                    "required": [
                        "question_col", "answer_col", "header_row",
                        "data_start_row", "format_type", "choices_col", "remarks_col",
                    ],
                    "additionalProperties": False,
                },
                "reasoning": {
                    "type": "object",
                    "properties": {
                        "question_col": {"type": "string"},
                        "answer_col": {"type": "string"},
                        "header_row": {"type": "string"},
                        "data_start_row": {"type": "string"},
                        "format_type": {"type": "string"},
                        "choices_col": {"type": "string"},
                        "remarks_col": {"type": "string"},
                    },
                    "required": [
                        "question_col", "answer_col", "header_row",
                        "data_start_row", "format_type", "choices_col", "remarks_col",
                    ],
                    "additionalProperties": False,
                },
            },
            "required": [
                "question_col", "answer_col", "header_row",
                "data_start_row", "format_type", "choices_col", "remarks_col",
                "confidence", "reasoning",
            ],
            "additionalProperties": False,
        },
    },
}

_MAX_PREVIEW_ROWS = 20


def _index_to_col_letter(idx: int) -> str:
    """0-indexed column index to Excel column letter (A, B, ..., Z, AA, ...)."""
    result = ""
    idx += 1
    while idx > 0:
        idx, remainder = divmod(idx - 1, 26)
        result = chr(65 + remainder) + result
    return result


def extract_preview(
    file_path: Path,
    file_format: str,
    sheet_name: Optional[str] = None,
    table_index: int = -1,
) -> Dict[str, Any]:
    """ファイルの先頭N行をプレビュー用に抽出する. table_index=-1 で自動選択."""
    if file_format == "docx":
        return _extract_docx_preview(file_path, table_index)
    return _extract_xlsx_preview(file_path, sheet_name)


def _extract_xlsx_preview(
    file_path: Path,
    sheet_name: Optional[str] = None,
) -> Dict[str, Any]:
    from openpyxl import load_workbook

    wb = load_workbook(str(file_path), read_only=False, data_only=True)
    sheet_names = wb.sheetnames
    ws = wb[sheet_name] if sheet_name and sheet_name in sheet_names else wb.active

    # 結合セルの値マップを構築
    merged_values: Dict[str, Any] = {}
    for merge_range in ws.merged_cells.ranges:
        top_left = ws.cell(merge_range.min_row, merge_range.min_col).value
        for row in range(merge_range.min_row, merge_range.max_row + 1):
            for col in range(merge_range.min_col, merge_range.max_col + 1):
                if row != merge_range.min_row or col != merge_range.min_col:
                    merged_values[(row, col)] = top_left

    max_col = ws.max_column or 1
    col_letters = [_index_to_col_letter(i) for i in range(max_col)]

    rows: List[Dict[str, Any]] = []
    total_rows = ws.max_row or 0
    for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=min(_MAX_PREVIEW_ROWS, total_rows), values_only=False), start=1):
        cells: List[str] = []
        for cell in row:
            val = cell.value
            if val is None:
                val = merged_values.get((cell.row, cell.column), None)
            cells.append(str(val) if val is not None else "")
        rows.append({"row_num": row_idx, "cells": cells})

    wb.close()
    return {
        "col_letters": col_letters,
        "rows": rows,
        "total_rows": total_rows,
        "sheet_names": sheet_names,
        "table_count": 0,
    }


def _extract_docx_preview(
    file_path: Path,
    table_index: int = -1,
) -> Dict[str, Any]:
    from docx import Document

    doc = Document(str(file_path))
    table_count = len(doc.tables)
    if table_count == 0:
        return {
            "col_letters": [],
            "rows": [],
            "total_rows": 0,
            "sheet_names": [],
            "table_count": 0,
            "tables_summary": [],
            "best_table_index": 0,
        }

    # 全テーブルのサマリを作成
    tables_summary: List[Dict[str, Any]] = []
    best_idx = 0
    best_size = 0
    for i, t in enumerate(doc.tables):
        t_rows = len(t.rows)
        t_cols = max((len(r.cells) for r in t.rows), default=0)
        header = ""
        if t.rows:
            header_cells = [c.text.strip()[:30] for c in t.rows[0].cells[:5]]
            header = " | ".join(header_cells)
        tables_summary.append({
            "index": i,
            "rows": t_rows,
            "cols": t_cols,
            "header": header,
        })
        size = t_rows * t_cols
        if size > best_size:
            best_size = size
            best_idx = i

    # table_index=-1 の場合は自動選択
    selected_idx = best_idx if table_index < 0 else min(table_index, table_count - 1)

    tbl = doc.tables[selected_idx]
    max_col = max((len(r.cells) for r in tbl.rows), default=0)
    col_letters = [_index_to_col_letter(i) for i in range(max_col)]

    rows: List[Dict[str, Any]] = []
    for row_idx, row in enumerate(tbl.rows[:_MAX_PREVIEW_ROWS], start=1):
        cells = [cell.text.strip() for cell in row.cells]
        cells.extend([""] * (max_col - len(cells)))
        rows.append({"row_num": row_idx, "cells": cells})

    header_texts: List[str] = []
    if tbl.rows:
        first_row = tbl.rows[0]
        header_texts = [cell.text.strip() for cell in first_row.cells]
        header_texts.extend([""] * (max_col - len(header_texts)))

    return {
        "col_letters": col_letters,
        "rows": rows,
        "total_rows": len(tbl.rows),
        "sheet_names": [],
        "table_count": table_count,
        "tables_summary": tables_summary,
        "best_table_index": best_idx,
        "selected_table_index": selected_idx,
        "header_texts": header_texts,
    }


def _build_user_prompt(preview: Dict[str, Any]) -> str:
    """プレビューデータからLLM用プロンプトを構築する."""
    col_letters = preview["col_letters"]
    lines = [f"以下はファイルの先頭{len(preview['rows'])}行のデータです。"]
    lines.append(f"総行数: {preview['total_rows']}")
    lines.append("")

    for row_data in preview["rows"]:
        row_num = row_data["row_num"]
        cells = row_data["cells"]
        parts = []
        for i, val in enumerate(cells):
            col = col_letters[i] if i < len(col_letters) else f"col{i}"
            display = val if val else "(空)"
            parts.append(f'{col}="{display}"')
        lines.append(f"行{row_num}: {' | '.join(parts)}")

    lines.append("")
    lines.append("このファイルの構造を分析し、列マッピングを提案してください。")
    return "\n".join(lines)


def _is_openai_model(model: str) -> bool:
    model_lower = model.lower()
    return any(prefix in model_lower for prefix in ("gpt-", "o1-", "o3-", "openai/"))


def analyze_format(
    preview: Dict[str, Any],
    model: str,
    api_key: str,
    user_hint: str = "",
) -> Dict[str, Any]:
    """プレビューデータをLLMに送り、フォーマット解析結果を返す."""
    user_prompt = _build_user_prompt(preview)
    if user_hint:
        user_prompt += f"\n\nユーザからの補足情報:\n{user_hint}"

    kwargs: dict = dict(
        model=model,
        max_tokens=4096,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        api_key=api_key or None,
    )
    if supports_response_schema(model, None):
        kwargs["response_format"] = _RESPONSE_SCHEMA
    elif _is_openai_model(model):
        kwargs["response_format"] = {"type": "json_object"}

    response = litellm.completion(**kwargs)
    response_text = response.choices[0].message.content

    if not response_text:
        logger.warning("Empty response from format analysis LLM")
        return _default_suggestion()

    try:
        result = json.loads(response_text)
    except json.JSONDecodeError:
        from .reviewer import _extract_json_object
        extracted = _extract_json_object(response_text)
        try:
            result = json.loads(extracted)
        except json.JSONDecodeError:
            logger.warning("Failed to parse format analysis response: %s", response_text[:200])
            return _default_suggestion()

    # 必須フィールドの存在確認
    for key in ("question_col", "answer_col", "header_row", "data_start_row"):
        if key not in result:
            logger.warning("Missing key %s in analysis result", key)
            return _default_suggestion()

    # デフォルト補完
    result.setdefault("format_type", "freetext")
    result.setdefault("choices_col", "")
    result.setdefault("remarks_col", "")
    result.setdefault("confidence", {})
    result.setdefault("reasoning", {})

    return result


def _default_suggestion() -> Dict[str, Any]:
    """解析失敗時のデフォルト提案."""
    return {
        "question_col": "D",
        "answer_col": "E",
        "header_row": 1,
        "data_start_row": 2,
        "format_type": "freetext",
        "choices_col": "",
        "remarks_col": "",
        "confidence": {
            "question_col": "low",
            "answer_col": "low",
            "header_row": "low",
            "data_start_row": "low",
            "format_type": "low",
            "choices_col": "low",
            "remarks_col": "low",
        },
        "reasoning": {
            "question_col": "自動解析に失敗したためデフォルト値です",
            "answer_col": "自動解析に失敗したためデフォルト値です",
            "header_row": "自動解析に失敗したためデフォルト値です",
            "data_start_row": "自動解析に失敗したためデフォルト値です",
            "format_type": "自動解析に失敗したためデフォルト値です",
            "choices_col": "自動解析に失敗したためデフォルト値です",
            "remarks_col": "自動解析に失敗したためデフォルト値です",
        },
    }


COLUMN_ROLES_READ = [
    "question",   # 質問・確認事項（システムが参照）
    "category",   # 分類・カテゴリ
    "number",     # 番号
    "reference",  # 参照情報（判断基準・エビデンス・設定例等）
    "remarks",    # 備考
]

COLUMN_ROLES_WRITE = [
    "answer",     # 回答欄（テキスト回答）（システムが記入）
    "judgment",   # 判定欄（○/△/×等の記号）
]

COLUMN_ROLES = COLUMN_ROLES_READ + COLUMN_ROLES_WRITE + ["other"]

_COLUMN_DEFS_SYSTEM_PROMPT = """\
あなたは金融機関のアンケート（質問票）の構造を分析する専門家です。
ヘッダー行とデータ行が提供されます。ワークフローに重要な列を特定し、各列の役割と説明を記述してください。

役割タグ（role）— 読み取り/書き込みで区分:

読み取り列（システムが参照する列）:
- question: 質問文・確認事項
- category: 分類・カテゴリ
- number: 通し番号
- reference: 参照情報（判断基準・エビデンス・設定例等）
- remarks: 備考

書き込み列（システムが回答を記入する列）:
- answer: 回答欄（テキスト回答）
- judgment: 判定欄（○/△/×等の記号）

その他:
- other: その他の重要列

重要でない列（空列、繰り返しの結合セル等）は含めなくて構いません。
必ず question と answer の役割を1つずつ含めてください。

注意事項:
- データ行の外側（ヘッダーより上、本体テーブルの右側/下側）に凡例テーブルや注釈がある場合、それらの列は answer や judgment に分類しないでください。
- answer 列は通常1〜2列です。3列以上を answer に分類するのは稀です。
- ヘッダーが2行にまたがる場合（結合セルのため）、両方の行を考慮して列の意味を判断してください。

また、行の構造を自然言語で説明してください。
この説明は、回答生成AIに「このアンケートの各行がどういう構造か」を伝えるために使います。

また、データ行の中に「見出し行」（セクション区切り、大分類の見出しなど、質問ではない行）が
あるかどうかを判定してください。ある場合、見出し行の識別方法を heading_pattern フィールドに記述してください。
見出し行がない場合は heading_pattern を空文字 "" にしてください（「存在しない」等のテキストは書かないでください）。

以下のJSON形式で回答してください:
{
  "column_definitions": [
    {"col": "A", "role": "number", "description": "通し番号"},
    {"col": "E", "role": "question", "description": "規定内容・確認事項"}
  ],
  "row_structure": "各行は1つの確認項目に対応。A列に通番、B-C列に分類、E列に質問、F列に回答を記入。",
  "heading_pattern": "番号列(A列)が空で、質問列にのみテキストがある行は見出し行"
}
"""


def _build_column_defs_prompt(
    preview: Dict[str, Any],
    header_row: int,
    data_start_row: int,
    user_hint: str = "",
) -> str:
    """2次分析用プロンプト: ヘッダー行＋データ行のみを送る."""
    col_letters = preview["col_letters"]
    lines = []

    # ヘッダー行（header_row〜data_start_row-1の全行を送信）
    for h_row_num in range(header_row, data_start_row):
        header = next((r for r in preview["rows"] if r["row_num"] == h_row_num), None)
        if header:
            parts = []
            for i, val in enumerate(header["cells"]):
                col = col_letters[i] if i < len(col_letters) else f"col{i}"
                parts.append(f'{col}="{val if val else "(空)"}"')
            lines.append(f"ヘッダー行（行{h_row_num}）: {' | '.join(parts)}")

    # ヘッダーより上の行があれば凡例/メタデータとして送信
    meta_rows = [r for r in preview["rows"] if r["row_num"] < header_row]
    if meta_rows:
        lines.append("\nヘッダーより上の行（メタデータ・凡例等。データ行ではありません）:")
        for row_data in meta_rows[:3]:
            parts = []
            for i, val in enumerate(row_data["cells"]):
                col = col_letters[i] if i < len(col_letters) else f"col{i}"
                parts.append(f'{col}="{val if val else "(空)"}"')
            lines.append(f"  行{row_data['row_num']}: {' | '.join(parts)}")
    lines.append("")

    data_rows = [r for r in preview["rows"] if r["row_num"] >= data_start_row][:5]
    lines.append("データ行:")
    for row_data in data_rows:
        parts = []
        for i, val in enumerate(row_data["cells"]):
            col = col_letters[i] if i < len(col_letters) else f"col{i}"
            display = val if val else "(空)"
            parts.append(f'{col}="{display}"')
        lines.append(f"  行{row_data['row_num']}: {' | '.join(parts)}")

    if user_hint:
        lines.append(f"\nユーザからの補足情報:\n{user_hint}")

    lines.append("\n重要な列の役割を特定し、行の構造を説明してください。")
    return "\n".join(lines)


def analyze_columns(
    preview: Dict[str, Any],
    header_row: int,
    data_start_row: int,
    model: str,
    api_key: str,
    user_hint: str = "",
    **kwargs: Any,
) -> Dict[str, Any]:
    """確定済みの基本設定を元に、列定義＋行構造をLLMで分析する（2次分析）."""
    user_prompt = _build_column_defs_prompt(preview, header_row, data_start_row, user_hint)

    llm_kwargs: dict = dict(
        model=model,
        max_tokens=4096,
        messages=[
            {"role": "system", "content": _COLUMN_DEFS_SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        api_key=api_key or None,
    )
    if _is_openai_model(model):
        llm_kwargs["response_format"] = {"type": "json_object"}

    response = litellm.completion(**llm_kwargs)
    response_text = response.choices[0].message.content

    if not response_text:
        logger.warning("Empty response from column definitions LLM")
        return _default_column_defs_suggestion(preview)

    try:
        result = json.loads(response_text)
    except json.JSONDecodeError:
        from .reviewer import _extract_json_object
        extracted = _extract_json_object(response_text)
        try:
            result = json.loads(extracted)
        except json.JSONDecodeError:
            logger.warning("Failed to parse column defs response: %s", response_text[:200])
            return _default_column_defs_suggestion(preview)

    if "column_definitions" not in result or not isinstance(result["column_definitions"], list):
        return _default_column_defs_suggestion(preview)

    result.setdefault("row_structure", "")
    result.setdefault("heading_pattern", "")

    # answer列数の上限バリデーション（3列以上は超過分を other に降格）
    answer_count = sum(1 for d in result["column_definitions"] if d.get("role") == "answer")
    if answer_count > 2:
        logger.warning("LLM assigned %d answer columns, capping to 2", answer_count)
        seen = 0
        for d in result["column_definitions"]:
            if d.get("role") == "answer":
                seen += 1
                if seen > 2:
                    d["role"] = "other"

    return result


def _default_column_defs_suggestion(preview: Dict[str, Any]) -> Dict[str, Any]:
    cols = preview.get("col_letters", [])
    defs = []
    if len(cols) >= 4:
        defs.append({"col": cols[3], "role": "question", "description": "質問・確認事項"})
    if len(cols) >= 5:
        defs.append({"col": cols[4], "role": "answer", "description": "回答欄"})
    return {
        "column_definitions": defs,
        "row_structure": "",
    }
