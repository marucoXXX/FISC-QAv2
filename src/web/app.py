"""FastAPI application for FISC-QAv2."""

from __future__ import annotations

import json
import os
import shutil
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any, List, Optional

from fastapi import APIRouter, FastAPI, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel

from ..config import Config
from . import db
from . import pipeline

UPLOAD_DIR = Path("data/uploads")



# --- Pydantic Models ---

class ReviewRequest(BaseModel):
    status: str
    comment: str = ""

class BulkReviewRequest(BaseModel):
    status: str
    question_nos: Optional[List[int]] = None

class BankCreate(BaseModel):
    name: str
    code: str
    notes: str = ""

class BankUpdate(BaseModel):
    name: Optional[str] = None
    code: Optional[str] = None
    notes: Optional[str] = None

class BankQaFileCreate(BaseModel):
    qa_file_name: str
    file_format: str = "xlsx"
    question_col: str = "D"
    answer_col: str = "E"
    header_row: int = 1
    data_start_row: int = 2
    table_index: int = 0
    format_type: str = "freetext"
    choices_col: str = ""
    remarks_col: str = ""

class BankQaFileUpdate(BaseModel):
    qa_file_name: Optional[str] = None
    file_format: Optional[str] = None
    question_col: Optional[str] = None
    answer_col: Optional[str] = None
    header_row: Optional[int] = None
    data_start_row: Optional[int] = None
    table_index: Optional[int] = None
    format_type: Optional[str] = None
    choices_col: Optional[str] = None
    remarks_col: Optional[str] = None

class CommonAnswerCreate(BaseModel):
    question_pattern: str
    answer_text: str
    category: str = ""
    source: str = ""

class CommonAnswerUpdate(BaseModel):
    question_pattern: Optional[str] = None
    answer_text: Optional[str] = None
    category: Optional[str] = None
    source: Optional[str] = None

class ConfirmItem(BaseModel):
    question_no: int
    confirmed: bool

class QuestionEdit(BaseModel):
    answer_text: str
    add_to_common: bool = False

class ConfigUpdate(BaseModel):
    kb_dir: Optional[str] = None
    model: Optional[str] = None
    token_budget: Optional[int] = None
    output_dir: Optional[str] = None

class KbFolderCreate(BaseModel):
    path: str
    label: str = ""

class KbFolderUpdate(BaseModel):
    path: Optional[str] = None
    label: Optional[str] = None


def create_router(
    db_path: Path,
    config: Config,
) -> APIRouter:
    """Create an APIRouter with all business API routes."""
    router = APIRouter()

    # --- Bank API ---

    @router.get("/api/banks")
    def list_banks() -> list[dict]:
        return db.list_banks(db_path)

    @router.post("/api/banks")
    def create_bank(req: BankCreate) -> dict:
        try:
            bank_id = db.create_bank(
                db_path, req.name, req.code, req.notes,
            )
        except Exception as e:
            if "UNIQUE" in str(e):
                raise HTTPException(409, "銀行名またはコードが重複しています")
            raise
        return {"id": bank_id}

    @router.get("/api/banks/{bank_id}")
    def get_bank(bank_id: int) -> dict:
        bank = db.get_bank(db_path, bank_id)
        if not bank:
            raise HTTPException(404, "銀行が見つかりません")
        return bank

    @router.put("/api/banks/{bank_id}")
    def update_bank(bank_id: int, req: BankUpdate) -> dict:
        ok = db.update_bank(db_path, bank_id, **req.model_dump(exclude_none=True))
        if not ok:
            raise HTTPException(404)
        return {"ok": True}

    @router.delete("/api/banks/{bank_id}")
    def delete_bank(bank_id: int) -> dict:
        ok = db.delete_bank(db_path, bank_id)
        if not ok:
            raise HTTPException(404)
        return {"ok": True}

    # --- Bank QA Files API ---

    @router.get("/api/banks/{bank_id}/qa-files")
    def list_bank_qa_files(bank_id: int) -> list[dict]:
        return db.list_bank_qa_files(db_path, bank_id)

    @router.post("/api/banks/{bank_id}/qa-files")
    def create_bank_qa_file(bank_id: int, req: BankQaFileCreate) -> dict:
        bank = db.get_bank(db_path, bank_id)
        if not bank:
            raise HTTPException(404, "銀行が見つかりません")
        try:
            qa_file_id = db.create_bank_qa_file(
                db_path, bank_id, req.qa_file_name, req.file_format,
                req.question_col, req.answer_col, req.header_row,
                req.data_start_row, req.table_index,
                req.format_type, req.choices_col, req.remarks_col,
            )
        except Exception as e:
            if "UNIQUE" in str(e):
                raise HTTPException(409, "同じ銀行に同名のQAファイルが既に存在します")
            raise
        return {"id": qa_file_id}

    @router.get("/api/banks/{bank_id}/qa-files/{qa_file_id}")
    def get_bank_qa_file(bank_id: int, qa_file_id: int) -> dict:
        qf = db.get_bank_qa_file(db_path, qa_file_id)
        if not qf:
            raise HTTPException(404)
        return qf

    @router.put("/api/banks/{bank_id}/qa-files/{qa_file_id}")
    def update_bank_qa_file(bank_id: int, qa_file_id: int, req: BankQaFileUpdate) -> dict:
        ok = db.update_bank_qa_file(db_path, qa_file_id, **req.model_dump(exclude_none=True))
        if not ok:
            raise HTTPException(404)
        return {"ok": True}

    @router.delete("/api/banks/{bank_id}/qa-files/{qa_file_id}")
    def delete_bank_qa_file(bank_id: int, qa_file_id: int) -> dict:
        ok = db.delete_bank_qa_file(db_path, qa_file_id)
        if not ok:
            raise HTTPException(404)
        return {"ok": True}

    # --- Past QA API ---

    @router.get("/api/banks/{bank_id}/past-answers")
    def list_past_answers(bank_id: int) -> list[dict]:
        return db.list_past_qa(db_path, bank_id)

    @router.post("/api/banks/{bank_id}/past-answers/upload")
    async def upload_past_answers(
        bank_id: int,
        file: UploadFile,
        qa_file_id: Optional[int] = Query(None),
    ) -> dict:
        bank = db.get_bank(db_path, bank_id)
        if not bank:
            raise HTTPException(404, "銀行が見つかりません")

        qf = None
        if qa_file_id:
            qf = db.get_bank_qa_file(db_path, qa_file_id)
            if not qf:
                raise HTTPException(404, "QAファイルが見つかりません")

        if not file.filename:
            raise HTTPException(400, "ファイル名がありません")

        content = await file.read()
        qa_pairs = _extract_qa_pairs(content, file.filename, qf)
        if not qa_pairs:
            raise HTTPException(400, "Q&Aペアを抽出できませんでした")

        count = db.bulk_add_past_qa(db_path, bank_id, qa_pairs, source_file=file.filename)
        return {"count": count, "message": f"{count}件のQ&Aペアを登録しました"}

    @router.delete("/api/banks/{bank_id}/past-answers/{past_qa_id}")
    def delete_past_answer(bank_id: int, past_qa_id: int) -> dict:
        ok = db.delete_past_qa(db_path, past_qa_id)
        if not ok:
            raise HTTPException(404)
        return {"ok": True}

    # --- Common Answers API ---

    @router.get("/api/common-answers")
    def list_common_answers(
        category: Optional[str] = Query(None),
        search: Optional[str] = Query(None),
    ) -> list[dict]:
        return db.list_common_answers(db_path, category=category, search=search)

    @router.post("/api/common-answers")
    def create_common_answer(req: CommonAnswerCreate) -> dict:
        common_id = db.create_common_answer(
            db_path, req.question_pattern, req.answer_text,
            req.category, req.source,
        )
        return {"id": common_id}

    @router.get("/api/common-answers/{common_id}")
    def get_common_answer(common_id: int) -> dict:
        ca = db.get_common_answer(db_path, common_id)
        if not ca:
            raise HTTPException(404)
        return ca

    @router.put("/api/common-answers/{common_id}")
    def update_common_answer(common_id: int, req: CommonAnswerUpdate) -> dict:
        ok = db.update_common_answer(db_path, common_id, **req.model_dump(exclude_none=True))
        if not ok:
            raise HTTPException(404)
        return {"ok": True}

    @router.delete("/api/common-answers/{common_id}")
    def delete_common_answer(common_id: int) -> dict:
        ok = db.delete_common_answer(db_path, common_id)
        if not ok:
            raise HTTPException(404)
        return {"ok": True}

    # --- Session API ---

    @router.get("/api/sessions")
    def list_sessions() -> list[dict]:
        return db.list_sessions(db_path)

    @router.post("/api/sessions")
    async def create_session(
        bank_id: int = Query(...),
        qa_file_id: Optional[int] = Query(None),
        file: UploadFile = ...,
    ) -> dict:
        bank = db.get_bank(db_path, bank_id)
        if not bank:
            raise HTTPException(404, "銀行が見つかりません")

        # Get format config from QA file
        qf = None
        if qa_file_id:
            qf = db.get_bank_qa_file(db_path, qa_file_id)
            if not qf:
                raise HTTPException(404, "QAファイルが見つかりません")

        if not file.filename:
            raise HTTPException(400, "ファイル名がありません")

        content = await file.read()

        # Save uploaded file
        session_id = db.create_session(
            db_path, bank_id, file.filename,
            source_file_name=file.filename,
            qa_file_id=qa_file_id,
        )
        upload_dir = UPLOAD_DIR / str(session_id)
        upload_dir.mkdir(parents=True, exist_ok=True)
        stored_path = upload_dir / file.filename
        stored_path.write_bytes(content)
        db.update_session(db_path, session_id, **{})  # trigger updated_at
        # Update stored path
        with db.get_conn(db_path) as conn:
            conn.execute(
                "UPDATE sessions SET source_file_path = ? WHERE id = ?",
                (str(stored_path), session_id),
            )

        # Extract questions from file
        from ..file_io import FormatConfig, read_questionnaire

        fc = FormatConfig(
            file_format=qf["file_format"] if qf else "xlsx",
            question_col=qf["question_col"] if qf else "D",
            answer_col=qf["answer_col"] if qf else "E",
            header_row=qf["header_row"] if qf else 1,
            data_start_row=qf["data_start_row"] if qf else 2,
            table_index=qf["table_index"] if qf else 0,
            format_type=qf["format_type"] if qf else "freetext",
            choices_col=qf["choices_col"] if qf else "",
            remarks_col=qf["remarks_col"] if qf else "",
        )
        questions = read_questionnaire(stored_path, fc)
        if not questions:
            raise HTTPException(400, "質問を抽出できませんでした")

        db.bulk_add_session_questions(db_path, session_id, [
            {"question_no": q.question_no, "major": q.major,
             "minor": q.minor, "question_text": q.question_text,
             "choices_text": q.choices_text, "remarks_text": q.remarks_text}
            for q in questions
        ])

        return {"session_id": session_id, "question_count": len(questions)}

    @router.get("/api/sessions/{session_id}")
    def get_session(session_id: int) -> dict:
        s = db.get_session(db_path, session_id)
        if not s:
            raise HTTPException(404)
        return s

    @router.delete("/api/sessions/{session_id}")
    def delete_session(session_id: int) -> dict:
        ok = db.delete_session(db_path, session_id)
        if not ok:
            raise HTTPException(404)
        return {"ok": True}

    @router.get("/api/sessions/{session_id}/questions")
    def get_session_questions(session_id: int) -> list[dict]:
        return db.get_session_questions(db_path, session_id)

    # Step2: Past answer matching
    @router.post("/api/sessions/{session_id}/step2/match")
    def step2_match(session_id: int, match_strategy: str = Query("cosine")) -> dict:
        session = db.get_session(db_path, session_id)
        if not session:
            raise HTTPException(404)

        questions = db.get_session_questions(db_path, session_id)
        past_qas = db.list_past_qa(db_path, session["bank_id"])

        q_dicts = [{"question_no": q["question_no"], "question_text": q["question_text"]} for q in questions]

        if match_strategy == "llm":
            from ..matcher import match_past_answers_llm
            matches = match_past_answers_llm(q_dicts, past_qas, api_key=config.api_key, model=config.model)
        elif match_strategy == "hybrid":
            from ..matcher import match_past_answers_hybrid
            matches = match_past_answers_hybrid(q_dicts, past_qas, api_key=config.api_key, model=config.model)
        else:
            from ..matcher import match_past_answers
            matches = match_past_answers(q_dicts, past_qas)

        for q_no, m in matches.items():
            db.update_session_question(db_path, session_id, q_no,
                matched_past_qa_id=m.matched_id,
                past_question_text=m.matched_question,
                past_answer_text=m.matched_answer,
                match_judgment=m.judgment,
                match_reason=m.reason,
            )

        return {"matched": len(matches), "total": len(questions)}

    @router.get("/api/sessions/{session_id}/step2/results")
    def step2_results(session_id: int) -> list[dict]:
        return db.get_session_questions(db_path, session_id)

    @router.put("/api/sessions/{session_id}/step2/confirm")
    def step2_confirm(session_id: int, items: List[ConfirmItem]) -> dict:
        count = 0
        for item in items:
            if item.confirmed:
                q = None
                for sq in db.get_session_questions(db_path, session_id):
                    if sq["question_no"] == item.question_no:
                        q = sq
                        break
                if q and q["past_answer_text"]:
                    db.update_session_question(db_path, session_id, item.question_no,
                        answer_source="past_match",
                        answer_text=q["past_answer_text"],
                        user_confirmed=1,
                        step_resolved=2,
                    )
                    count += 1
            else:
                db.update_session_question(db_path, session_id, item.question_no,
                    user_confirmed=-1,
                )
        db.update_session(db_path, session_id, current_step=3)
        return {"confirmed": count}

    # Step3: Common answer matching
    @router.post("/api/sessions/{session_id}/step3/match")
    def step3_match(session_id: int) -> dict:
        session = db.get_session(db_path, session_id)
        if not session:
            raise HTTPException(404)

        unresolved = db.get_unresolved_questions(db_path, session_id)
        common_list = db.list_common_answers(db_path)

        if not unresolved or not common_list:
            return {"matched": 0, "total": len(unresolved)}

        from ..matcher import match_common_answers
        matches = match_common_answers(
            [{"question_no": q["question_no"], "question_text": q["question_text"]} for q in unresolved],
            common_list,
            api_key=config.api_key,
            model=config.model,
        )

        for q_no, m in matches.items():
            db.update_session_question(db_path, session_id, q_no,
                matched_common_id=m.matched_id,
                common_answer_text=m.matched_answer,
            )

        return {"matched": len(matches), "total": len(unresolved)}

    @router.get("/api/sessions/{session_id}/step3/results")
    def step3_results(session_id: int) -> list[dict]:
        return [q for q in db.get_session_questions(db_path, session_id)
                if q["answer_source"] == "pending"]

    @router.put("/api/sessions/{session_id}/step3/confirm")
    def step3_confirm(session_id: int, items: List[ConfirmItem]) -> dict:
        count = 0
        for item in items:
            if item.confirmed:
                q = None
                for sq in db.get_session_questions(db_path, session_id):
                    if sq["question_no"] == item.question_no:
                        q = sq
                        break
                if q and q["common_answer_text"]:
                    db.update_session_question(db_path, session_id, item.question_no,
                        answer_source="common_match",
                        answer_text=q["common_answer_text"],
                        user_confirmed=1,
                        step_resolved=3,
                    )
                    count += 1
            else:
                db.update_session_question(db_path, session_id, item.question_no,
                    user_confirmed=-1,
                )
        db.update_session(db_path, session_id, current_step=4)
        return {"confirmed": count}

    # Step4: AI generation for unresolved questions
    @router.post("/api/sessions/{session_id}/step4/generate")
    def step4_generate(session_id: int) -> dict:
        session = db.get_session(db_path, session_id)
        if not session:
            raise HTTPException(404)

        unresolved = db.get_unresolved_questions(db_path, session_id)
        if not unresolved:
            db.update_session(db_path, session_id, current_step=5)
            return {"job_id": None, "message": "生成が必要な質問はありません", "skipped": True}

        job_id = pipeline.start_session_pipeline_job(
            session_id=session_id,
            unresolved_questions=unresolved,
            config=config,
            db_path=db_path,
        )
        return {"job_id": job_id}

    @router.get("/api/sessions/{session_id}/step4/progress/{job_id}")
    def step4_progress(session_id: int, job_id: str) -> StreamingResponse:
        job = pipeline.get_job(job_id)
        if not job:
            raise HTTPException(404)
        return StreamingResponse(
            pipeline.get_job_progress_sse(job_id),
            media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )

    @router.get("/api/sessions/{session_id}/step4/status/{job_id}")
    def step4_status(session_id: int, job_id: str) -> dict:
        job = pipeline.get_job(job_id)
        if not job:
            raise HTTPException(404)
        return {
            "job_id": job.job_id,
            "status": job.status,
            "progress": job.progress,
            "error": job.error,
        }

    # Step5: Final review
    @router.get("/api/sessions/{session_id}/step5/summary")
    def step5_summary(session_id: int) -> dict:
        session = db.get_session(db_path, session_id)
        if not session:
            raise HTTPException(404)
        questions = db.get_session_questions(db_path, session_id)
        stats = {
            "total": len(questions),
            "past_match": sum(1 for q in questions if q["answer_source"] == "past_match"),
            "common_match": sum(1 for q in questions if q["answer_source"] == "common_match"),
            "generated": sum(1 for q in questions if q["answer_source"] == "generated"),
            "manual": sum(1 for q in questions if q["answer_source"] == "manual"),
            "pending": sum(1 for q in questions if q["answer_source"] == "pending"),
        }
        return {"session": session, "questions": questions, "stats": stats}

    @router.put("/api/sessions/{session_id}/questions/{question_no}")
    def edit_question_answer(session_id: int, question_no: int, req: QuestionEdit) -> dict:
        updates = {"answer_text": req.answer_text, "add_to_common": 1 if req.add_to_common else 0}
        q = None
        for sq in db.get_session_questions(db_path, session_id):
            if sq["question_no"] == question_no:
                q = sq
                break
        if q and q["answer_source"] == "pending":
            updates["answer_source"] = "manual"
            updates["step_resolved"] = 5
        ok = db.update_session_question(db_path, session_id, question_no, **updates)
        if not ok:
            raise HTTPException(404)
        return {"ok": True}

    @router.put("/api/sessions/{session_id}/step5/finalize")
    def step5_finalize(session_id: int) -> dict:
        session = db.get_session(db_path, session_id)
        if not session:
            raise HTTPException(404)

        questions = db.get_session_questions(db_path, session_id)

        # Auto-accumulate to past_qa
        for q in questions:
            if q["answer_text"]:
                db.add_past_qa(
                    db_path, session["bank_id"],
                    q["question_text"], q["answer_text"],
                    source_file=session.get("source_file_name", ""),
                )

        # Add selected to common_answers
        for q in questions:
            if q["add_to_common"] and q["answer_text"]:
                db.create_common_answer(
                    db_path, q["question_text"], q["answer_text"],
                    category=q.get("major", ""),
                )

        db.update_session(db_path, session_id, status="completed", current_step=5)
        return {"ok": True, "accumulated": len([q for q in questions if q["answer_text"]])}

    @router.get("/api/sessions/{session_id}/export")
    def export_session(session_id: int) -> StreamingResponse:
        session = db.get_session(db_path, session_id)
        if not session:
            raise HTTPException(404)

        questions = db.get_session_questions(db_path, session_id)
        answers = {q["question_no"]: q["answer_text"] for q in questions if q["answer_text"]}

        # assessment型の場合、○/△/×をchoices_colに書き込む
        choices = None
        fmt_type = session.get("format_type") or "freetext"
        if fmt_type == "assessment" and session.get("choices_col"):
            choices = {q["question_no"]: q["assessment_mark"]
                       for q in questions if q.get("assessment_mark")}

        source_path = Path(session.get("source_file_path", ""))
        if source_path.exists():
            from ..file_io import FormatConfig, write_answers_to_original

            fc = FormatConfig(
                file_format=session["file_format"] or "xlsx",
                question_col=session["question_col"] or "D",
                answer_col=session["answer_col"] or "E",
                header_row=session["header_row"] or 1,
                data_start_row=session["data_start_row"] or 2,
                table_index=session["table_index"] or 0,
                format_type=fmt_type,
                choices_col=session.get("choices_col") or "",
                remarks_col=session.get("remarks_col") or "",
            )
            output_dir = Path(config.output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            ext = ".docx" if (session["file_format"] or "xlsx") == "docx" else ".xlsx"
            output_path = output_dir / f"FISC回答_{session_id}{ext}"
            write_answers_to_original(source_path, answers, fc, output_path, choices=choices)

            media = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                if ext == ".docx" else
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            safe_name = f"FISC_answer_{session_id}{ext}"
            return StreamingResponse(
                open(str(output_path), "rb"),
                media_type=media,
                headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
            )

        # Fallback: generate new Excel
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "FISC回答"
        ws.append(["No.", "大分類", "小分類", "質問", "回答", "ソース", "確信度"])
        for q in questions:
            ws.append([
                q["question_no"], q["major"], q["minor"],
                q["question_text"], q["answer_text"],
                q["answer_source"], q["confidence"],
            ])
        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="FISC_answer_{session_id}.xlsx"'},
        )

    # --- Legacy Run API Routes ---

    @router.post("/api/runs/import")
    async def import_run(file: UploadFile) -> dict:
        if not file.filename:
            raise HTTPException(400, "ファイル名がありません")

        content = await file.read()

        if file.filename.endswith(".xlsx"):
            questions, answers, notes = _parse_excel(content)
        elif file.filename.endswith(".json"):
            questions, answers, notes = _parse_json(content)
        else:
            raise HTTPException(400, "Excel (.xlsx) または JSON (.json) ファイルのみ対応")

        run_id = db.import_run(
            db_path, file.filename, questions, answers, notes, source_file=file.filename,
        )
        return {"run_id": run_id, "message": f"インポート完了: {len(answers)}件の回答"}

    @router.get("/api/runs")
    def list_runs() -> list[dict]:
        return db.list_runs(db_path)

    @router.get("/api/runs/{run_id}")
    def get_run(run_id: int) -> dict:
        run = db.get_run(db_path, run_id)
        if not run:
            raise HTTPException(404, "実行結果が見つかりません")
        return run

    @router.get("/api/runs/{run_id}/stats")
    def get_stats(run_id: int) -> dict:
        run = db.get_run(db_path, run_id)
        if not run:
            raise HTTPException(404)
        return db.get_run_stats(db_path, run_id)

    @router.get("/api/runs/{run_id}/answers")
    def get_answers(
        run_id: int,
        review_status: Optional[str] = Query(None),
    ) -> list[dict]:
        return db.get_answers(db_path, run_id, review_status=review_status)

    @router.get("/api/runs/{run_id}/answers/{question_no}")
    def get_answer(run_id: int, question_no: int) -> dict:
        ans = db.get_answer(db_path, run_id, question_no)
        if not ans:
            raise HTTPException(404, "回答が見つかりません")
        notes = [
            n for n in db.get_review_notes(db_path, run_id)
            if n["question_no"] == question_no
        ]
        ans["review_notes"] = notes
        return ans

    @router.put("/api/runs/{run_id}/answers/{question_no}/review")
    def set_review(run_id: int, question_no: int, req: ReviewRequest) -> dict:
        if req.status not in ("approved", "rejected", "needs_revision", "pending"):
            raise HTTPException(400, "無効なステータス")
        ok = db.set_review(db_path, run_id, question_no, req.status, req.comment)
        if not ok:
            raise HTTPException(404)
        return {"ok": True}

    @router.put("/api/runs/{run_id}/bulk-review")
    def bulk_review(run_id: int, req: BulkReviewRequest) -> dict:
        if req.status not in ("approved", "rejected", "needs_revision", "pending"):
            raise HTTPException(400, "無効なステータス")
        count = db.bulk_set_review(db_path, run_id, req.status, req.question_nos)
        return {"ok": True, "updated": count}

    @router.get("/api/runs/{run_id}/review-notes")
    def get_review_notes(run_id: int) -> list[dict]:
        return db.get_review_notes(db_path, run_id)

    @router.get("/api/runs/{run_id}/export")
    def export_approved(run_id: int) -> StreamingResponse:
        run = db.get_run(db_path, run_id)
        if not run:
            raise HTTPException(404)
        answers = db.get_answers(db_path, run_id, review_status="approved")
        all_answers = db.get_answers(db_path, run_id)

        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "FISC回答結果（確定）"
        ws.append(["No.", "大分類", "小分類", "質問内容", "回答", "対応状況",
                    "根拠ソース", "確信度", "レビュー", "備考"])

        approved_nos = {a["question_no"] for a in answers}
        for a in all_answers:
            ws.append([
                a["question_no"],
                a["major"],
                a["minor"],
                a["question"],
                a["answer"] if a["question_no"] in approved_nos else "",
                a["status"] if a["question_no"] in approved_nos else "未承認",
                ", ".join(a["source_references"]),
                a["confidence"],
                a["review_status"],
                a["flag"] or "",
            ])

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)

        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="FISC_approved_{run_id}.xlsx"'},
        )

    @router.delete("/api/runs/{run_id}")
    def delete_run(run_id: int) -> dict:
        with db.get_conn(db_path) as conn:
            conn.execute("DELETE FROM review_notes WHERE run_id = ?", (run_id,))
            conn.execute("DELETE FROM answers WHERE run_id = ?", (run_id,))
            cur = conn.execute("DELETE FROM runs WHERE id = ?", (run_id,))
            if cur.rowcount == 0:
                raise HTTPException(404)
        return {"ok": True}

    # --- Pipeline API ---

    @router.post("/api/pipeline/start")
    async def start_pipeline(file: UploadFile) -> dict:
        if not file.filename or not file.filename.endswith(".xlsx"):
            raise HTTPException(400, "質問票 Excel (.xlsx) ファイルが必要です")

        content = await file.read()
        tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        tmp.write(content)
        tmp.close()

        kb_folders = db.list_kb_folders(db_path)
        if kb_folders:
            kb_dirs = [Path(f["path"]) for f in kb_folders]
        else:
            kb_dirs = [Path(config.kb_dir)]

        existing_dirs = [d for d in kb_dirs if d.exists()]
        if not existing_dirs:
            raise HTTPException(400, f"KB ディレクトリが見つかりません: {kb_dirs}")

        job_id = pipeline.start_pipeline_job(
            questionnaire_path=Path(tmp.name),
            kb_dirs=existing_dirs,
            config=config,
            db_path=db_path,
        )
        return {"job_id": job_id}

    @router.get("/api/pipeline/{job_id}/progress")
    def pipeline_progress(job_id: str) -> StreamingResponse:
        job = pipeline.get_job(job_id)
        if not job:
            raise HTTPException(404, "ジョブが見つかりません")
        return StreamingResponse(
            pipeline.get_job_progress_sse(job_id),
            media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )

    @router.get("/api/pipeline/{job_id}/status")
    def pipeline_status(job_id: str) -> dict:
        job = pipeline.get_job(job_id)
        if not job:
            raise HTTPException(404, "ジョブが見つかりません")
        return {
            "job_id": job.job_id,
            "status": job.status,
            "progress": job.progress,
            "result_run_id": job.result_run_id,
            "error": job.error,
        }

    # --- KB Folders API ---

    @router.get("/api/kb-folders")
    def list_kb_folders() -> list[dict]:
        return db.list_kb_folders(db_path)

    @router.post("/api/kb-folders", status_code=201)
    def create_kb_folder(req: KbFolderCreate) -> dict:
        folder_id = db.create_kb_folder(db_path, path=req.path, label=req.label)
        return {"id": folder_id}

    @router.put("/api/kb-folders/{folder_id}")
    def update_kb_folder(folder_id: int, req: KbFolderUpdate) -> dict:
        updated = db.update_kb_folder(db_path, folder_id, **req.model_dump(exclude_none=True))
        if not updated:
            raise HTTPException(404, "KB folder not found")
        return {"ok": True}

    @router.delete("/api/kb-folders/{folder_id}")
    def delete_kb_folder(folder_id: int) -> dict:
        deleted = db.delete_kb_folder(db_path, folder_id)
        if not deleted:
            raise HTTPException(404, "KB folder not found")
        return {"ok": True}

    # --- Config API ---

    @router.get("/api/config")
    def get_config() -> dict:
        return {
            "kb_dir": config.kb_dir,
            "model": config.model,
            "token_budget": config.token_budget_per_reader,
            "output_dir": config.output_dir,
            "api_key_set": bool(config.api_key),
        }

    @router.put("/api/config")
    def update_config(req: ConfigUpdate) -> dict:
        if req.kb_dir is not None:
            config.kb_dir = req.kb_dir
        if req.model is not None:
            config.model = req.model
        if req.token_budget is not None:
            config.token_budget_per_reader = req.token_budget
        if req.output_dir is not None:
            config.output_dir = req.output_dir
        return {"ok": True}

    return router


def create_app(
    db_path: Path = db.DEFAULT_DB_PATH,
    config: Optional[Config] = None,
) -> FastAPI:
    config = config or Config()
    db.init_db(db_path)

    app = FastAPI(title="FISC-QAv2", version="0.2.0")
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["http://localhost:5173"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

    router = create_router(db_path, config)
    app.include_router(router)

    return app


def _parse_excel(content: bytes) -> tuple[list[dict], list[dict], list[dict]]:
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
    ws = wb.active
    questions = []
    answers = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue
        no = int(row[0])
        questions.append({
            "no": no,
            "major": str(row[1] or ""),
            "minor": str(row[2] or ""),
            "question": str(row[3] or ""),
        })
        answers.append({
            "question_no": no,
            "answer": str(row[4] or ""),
            "status": str(row[5] or ""),
            "source_references": [s.strip() for s in str(row[6] or "").split(",") if s.strip()],
            "confidence": str(row[7] or "low"),
            "key_excerpt": "",
            "flag": str(row[8] or "") or None,
        })

    notes = []
    if "レビュー指摘" in wb.sheetnames:
        ws2 = wb["レビュー指摘"]
        for row in ws2.iter_rows(min_row=2, values_only=True):
            if row[0] is None:
                continue
            notes.append({
                "question_no": int(row[0]),
                "issue_type": str(row[1] or "other"),
                "severity": str(row[2] or "medium"),
                "description": str(row[3] or ""),
                "suggestion": str(row[4] or "") if len(row) > 4 else "",
            })

    wb.close()
    return questions, answers, notes


def _parse_json(content: bytes) -> tuple[list[dict], list[dict], list[dict]]:
    data = json.loads(content.decode("utf-8"))

    questions = data.get("questions", [])
    answers = data.get("answers", [])
    notes = data.get("review_notes", [])

    return questions, answers, notes


def _extract_qa_pairs(content: bytes, filename: str, qf: dict | None) -> list[dict]:
    """QAファイルのフォーマット設定に基づいてファイルからQ&Aペアを抽出する"""
    q_col = qf["question_col"] if qf else "D"
    a_col = qf["answer_col"] if qf else "E"
    header_row = qf["header_row"] if qf else 1
    data_start = qf["data_start_row"] if qf else 2
    c_col = qf.get("choices_col", "") if qf else ""
    r_col = qf.get("remarks_col", "") if qf else ""

    if filename.endswith(".xlsx"):
        return _extract_qa_from_excel(content, q_col, a_col, header_row, data_start, c_col, r_col)
    elif filename.endswith(".docx"):
        table_index = qf["table_index"] if qf else 0
        return _extract_qa_from_docx(content, table_index, q_col, a_col, data_start, c_col, r_col)
    else:
        return []


def _col_letter_to_index(col: str) -> int:
    """Excel列名を0始まりインデックスに変換 (A=0, B=1, ...)"""
    result = 0
    for c in col.upper():
        result = result * 26 + (ord(c) - ord('A') + 1)
    return result - 1


def _extract_qa_from_excel(
    content: bytes, q_col: str, a_col: str, header_row: int, data_start: int,
    c_col: str = "", r_col: str = "",
) -> list[dict]:
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
    ws = wb.active
    q_idx = _col_letter_to_index(q_col)
    a_idx = _col_letter_to_index(a_col)
    c_idx = _col_letter_to_index(c_col) if c_col else -1
    r_idx = _col_letter_to_index(r_col) if r_col else -1

    pairs = []
    for row in ws.iter_rows(min_row=data_start, values_only=True):
        if len(row) <= max(q_idx, a_idx):
            continue
        q_text = str(row[q_idx] or "").strip()
        a_text = str(row[a_idx] or "").strip()
        if q_text and a_text:
            pair: dict = {"question_text": q_text, "answer_text": a_text}
            if 0 <= c_idx < len(row):
                pair["choices_text"] = str(row[c_idx] or "").strip()
            if 0 <= r_idx < len(row):
                pair["remarks_text"] = str(row[r_idx] or "").strip()
            pairs.append(pair)

    wb.close()
    return pairs


def _extract_qa_from_docx(
    content: bytes, table_index: int, q_col: str, a_col: str, data_start: int,
    c_col: str = "", r_col: str = "",
) -> list[dict]:
    from docx import Document

    doc = Document(BytesIO(content))
    if table_index >= len(doc.tables):
        return []

    table = doc.tables[table_index]
    q_idx = _col_letter_to_index(q_col)
    a_idx = _col_letter_to_index(a_col)
    c_idx = _col_letter_to_index(c_col) if c_col else -1
    r_idx = _col_letter_to_index(r_col) if r_col else -1

    pairs = []
    for i, row in enumerate(table.rows):
        if i < data_start - 1:
            continue
        cells = row.cells
        if len(cells) <= max(q_idx, a_idx):
            continue
        q_text = cells[q_idx].text.strip()
        a_text = cells[a_idx].text.strip()
        if q_text and a_text:
            pair: dict = {"question_text": q_text, "answer_text": a_text}
            if 0 <= c_idx < len(cells):
                pair["choices_text"] = cells[c_idx].text.strip()
            if 0 <= r_idx < len(cells):
                pair["remarks_text"] = cells[r_idx].text.strip()
            pairs.append(pair)

    return pairs
