"""File I/O module for reading/writing questionnaires in bank-specific formats."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any


@dataclass
class FormatConfig:
    file_format: str  # 'xlsx' or 'docx'
    question_col: str
    answer_col: str
    header_row: int
    data_start_row: int
    table_index: int = 0
    format_type: str = "freetext"  # 'choices' | 'assessment' | 'freetext'
    choices_col: str = ""
    remarks_col: str = ""


def derive_format_config(
    column_definitions: list,
    header_row: int,
    data_start_row: int,
    file_format: str = "xlsx",
    table_index: int = 0,
) -> FormatConfig:
    """column_definitions から FormatConfig を導出する."""
    def _find_col(role: str) -> str:
        return next((d["col"] for d in column_definitions if d.get("role") == role), "")

    q_col = _find_col("question") or "D"
    a_col = _find_col("answer") or "E"
    c_col = _find_col("judgment")
    r_col = _find_col("remarks")

    # format_type を判定列の有無から導出
    if c_col:
        fmt = "assessment"
    else:
        fmt = "freetext"

    return FormatConfig(
        file_format=file_format,
        question_col=q_col,
        answer_col=a_col,
        header_row=header_row,
        data_start_row=data_start_row,
        table_index=table_index,
        format_type=fmt,
        choices_col=c_col,
        remarks_col=r_col,
    )


_ROLES_READ = {"question", "category", "number", "reference", "remarks"}
_ROLES_WRITE = {"answer", "judgment"}

_ROLE_LABELS = {
    "question": "質問・確認事項",
    "answer": "回答欄",
    "category": "分類",
    "number": "番号",
    "reference": "参照情報",
    "remarks": "備考",
    "judgment": "判定欄",
    "other": "その他",
}


def build_format_context(column_defs: list, row_structure: str) -> str:
    """column_definitions + row_structure から Reader LLM 用のフォーマット文脈テキストを生成."""
    lines = []

    if row_structure:
        lines.append(f"【行構造】\n{row_structure}")

    read_cols = [d for d in column_defs if d.get("role") in _ROLES_READ]
    write_cols = [d for d in column_defs if d.get("role") in _ROLES_WRITE]

    if read_cols:
        lines.append("\n【読み取り列（参照情報）】")
        for d in read_cols:
            lines.append(f"  {d['col']}列: {d.get('description', '')} ({_ROLE_LABELS.get(d['role'], d['role'])})")

    if write_cols:
        lines.append("\n【書き込み列（回答を記入する列）】")
        for d in write_cols:
            role = d.get("role", "")
            desc = d.get("description", "")
            if role == "answer":
                lines.append(f"  {d['col']}列: {desc} → テキストで回答してください")
            elif role == "judgment":
                lines.append(f"  {d['col']}列: {desc} → ○/△/×の記号で判定してください")

    return "\n".join(lines)


@dataclass
class Question:
    question_no: int
    question_text: str
    major: str = ""
    minor: str = ""
    choices_text: str = ""
    remarks_text: str = ""
    extra_columns: dict = None  # type: ignore[assignment]
    is_heading: bool = False

    def __post_init__(self):
        if self.extra_columns is None:
            self.extra_columns = {}


def _col_letter_to_index(col: str) -> int:
    result = 0
    for c in col.upper():
        result = result * 26 + (ord(c) - ord('A') + 1)
    return result - 1


def _detect_heading(q: Question, heading_pattern: str, number_col_idx: int = -1, row_values: tuple = ()) -> bool:
    """見出し行かどうかを複合ルールで判定する。

    heading_pattern テキストに依存せず、構造ルールを常に適用:
    1. 全列同一値ルール: 行内の全セル（2列以上）が同じテキスト → 見出し行（結合セル）
    2. 【】ルール: テキストが【...】で囲まれている → 見出し行
    3. 番号列空ルール: number列が空で question列にテキストがある → 見出し行候補
    """
    import re
    text = q.question_text.strip()
    if not text:
        return False

    # ルール1: 全列同一値（結合セルの特徴）
    if len(row_values) >= 2:
        non_empty = [str(v or "").strip() for v in row_values if str(v or "").strip()]
        if non_empty and all(v == non_empty[0] for v in non_empty):
            return True

    # ルール2: 【】で囲まれたテキスト
    if re.match(r"^【.+】$", text):
        return True

    # ルール3: 番号列が空で質問列にテキストがある
    if number_col_idx >= 0 and number_col_idx < len(row_values):
        num_val = str(row_values[number_col_idx] or "").strip()
        if not num_val:
            # 他の列（質問列以外）がほぼ空ならば見出し行
            non_empty_count = sum(1 for v in row_values if str(v or "").strip())
            if non_empty_count <= 1:
                return True

    # heading_pattern テキストのキーワードマッチ（追加ルール）
    if heading_pattern:
        pat = heading_pattern.lower()
        if ("番号" in pat and "空" in pat) or ("number" in pat and "empty" in pat):
            if number_col_idx >= 0 and number_col_idx < len(row_values):
                num_val = str(row_values[number_col_idx] or "").strip()
                if not num_val and text:
                    return True

    return False


def read_questionnaire(
    path: Path,
    config: FormatConfig,
    column_definitions: list | None = None,
    heading_pattern: str = "",
) -> list[Question]:
    if config.file_format == "docx":
        return _read_docx_questionnaire(path, config, column_definitions, heading_pattern)
    return _read_xlsx_questionnaire(path, config, column_definitions, heading_pattern)


def _read_xlsx_questionnaire(
    path: Path,
    config: FormatConfig,
    column_definitions: list | None = None,
    heading_pattern: str = "",
) -> list[Question]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    ws = wb.active
    q_idx = _col_letter_to_index(config.question_col)
    c_idx = _col_letter_to_index(config.choices_col) if config.choices_col else -1
    r_idx = _col_letter_to_index(config.remarks_col) if config.remarks_col else -1

    # 読み取り列（question以外）のインデックスマップ
    extra_col_map: list[tuple[int, dict]] = []
    number_col_idx = -1
    if column_definitions:
        for d in column_definitions:
            if d.get("role") in _ROLES_READ and d.get("role") != "question":
                idx = _col_letter_to_index(d["col"])
                extra_col_map.append((idx, d))
            if d.get("role") == "number":
                number_col_idx = _col_letter_to_index(d["col"])

    questions = []
    for i, row in enumerate(ws.iter_rows(min_row=config.data_start_row, values_only=True), start=1):
        if len(row) <= q_idx:
            continue
        q_text = str(row[q_idx] or "").strip()
        if not q_text:
            continue
        major = str(row[1] or "").strip() if len(row) > 1 else ""
        minor = str(row[2] or "").strip() if len(row) > 2 else ""
        choices = str(row[c_idx] or "").strip() if 0 <= c_idx < len(row) else ""
        remarks = str(row[r_idx] or "").strip() if 0 <= r_idx < len(row) else ""

        extra = {}
        for idx, d in extra_col_map:
            val = str(row[idx] or "").strip() if idx < len(row) else ""
            if val:
                extra[d["col"]] = {
                    "role": d["role"],
                    "description": d.get("description", ""),
                    "value": val,
                }

        q = Question(
            question_no=i,
            question_text=q_text,
            major=major,
            minor=minor,
            choices_text=choices,
            remarks_text=remarks,
            extra_columns=extra,
        )
        if heading_pattern:
            q.is_heading = _detect_heading(q, heading_pattern, number_col_idx, row)
        questions.append(q)

    wb.close()
    return questions


def _read_docx_questionnaire(
    path: Path,
    config: FormatConfig,
    column_definitions: list | None = None,
    heading_pattern: str = "",
) -> list[Question]:
    from docx import Document

    doc = Document(str(path))
    if config.table_index >= len(doc.tables):
        return []

    table = doc.tables[config.table_index]
    q_idx = _col_letter_to_index(config.question_col)
    c_idx = _col_letter_to_index(config.choices_col) if config.choices_col else -1
    r_idx = _col_letter_to_index(config.remarks_col) if config.remarks_col else -1

    extra_col_map: list[tuple[int, dict]] = []
    number_col_idx = -1
    if column_definitions:
        for d in column_definitions:
            if d.get("role") in _ROLES_READ and d.get("role") != "question":
                idx = _col_letter_to_index(d["col"])
                extra_col_map.append((idx, d))
            if d.get("role") == "number":
                number_col_idx = _col_letter_to_index(d["col"])

    questions = []
    no = 0
    for i, row in enumerate(table.rows):
        if i < config.data_start_row - 1:
            continue
        cells = row.cells
        if len(cells) <= q_idx:
            continue
        q_text = cells[q_idx].text.strip()
        if not q_text:
            continue
        no += 1
        major = cells[1].text.strip() if len(cells) > 1 else ""
        minor = cells[2].text.strip() if len(cells) > 2 else ""
        choices = cells[c_idx].text.strip() if 0 <= c_idx < len(cells) else ""
        remarks = cells[r_idx].text.strip() if 0 <= r_idx < len(cells) else ""

        extra = {}
        for idx, d in extra_col_map:
            val = cells[idx].text.strip() if idx < len(cells) else ""
            if val:
                extra[d["col"]] = {
                    "role": d["role"],
                    "description": d.get("description", ""),
                    "value": val,
                }

        row_vals = tuple(c.text.strip() for c in cells)
        q = Question(
            question_no=no,
            question_text=q_text,
            major=major,
            minor=minor,
            choices_text=choices,
            remarks_text=remarks,
            extra_columns=extra,
        )
        if heading_pattern:
            q.is_heading = _detect_heading(q, heading_pattern, number_col_idx, row_vals)
        questions.append(q)

    return questions


def write_answers_to_original(
    original_path: Path,
    answers: dict[int, str],
    config: FormatConfig,
    output_path: Path,
    choices: dict[int, str] | None = None,
    column_definitions: list | None = None,
) -> Path:
    """元ファイルの回答列に書き込み、フォーマットを維持して出力

    choices: assessment型の場合、choices_col に書き込む値（例: {1: "○", 2: "△"}）
    column_definitions: 新形式の列定義。指定時は全書き込み列に対応。
    """
    # column_definitions から書き込み列を動的に決定
    write_cols: list[tuple[str, str]] = []  # [(col_letter, role), ...]
    if column_definitions:
        for d in column_definitions:
            if d.get("role") in _ROLES_WRITE:
                write_cols.append((d["col"], d["role"]))

    if config.file_format == "docx":
        return _write_docx_answers(original_path, answers, config, output_path, choices, write_cols)
    return _write_xlsx_answers(original_path, answers, config, output_path, choices, write_cols)


def _write_xlsx_answers(
    original_path: Path,
    answers: dict[int, str],
    config: FormatConfig,
    output_path: Path,
    choices: dict[int, str] | None = None,
    write_cols: list[tuple[str, str]] | None = None,
) -> Path:
    from openpyxl import load_workbook
    from openpyxl.cell.cell import MergedCell

    wb = load_workbook(str(original_path))
    ws = wb.active

    # 書き込み先の列を決定
    answer_cols: list[str] = []
    judgment_cols: list[str] = []
    if write_cols:
        for col_letter, role in write_cols:
            if role == "answer":
                answer_cols.append(col_letter.upper())
            elif role == "judgment":
                judgment_cols.append(col_letter.upper())
    else:
        # 後方互換: 旧フィールドを使用
        answer_cols = [config.answer_col.upper()]
        if config.choices_col:
            judgment_cols = [config.choices_col.upper()]

    no = 0
    for row_idx in range(config.data_start_row, ws.max_row + 1):
        q_idx = _col_letter_to_index(config.question_col)
        q_cell = ws.cell(row=row_idx, column=q_idx + 1)
        if isinstance(q_cell, MergedCell) or not q_cell.value or not str(q_cell.value).strip():
            continue
        no += 1
        # 回答列に書き込み
        if no in answers:
            for a_col in answer_cols:
                cell = ws[f"{a_col}{row_idx}"]
                if not isinstance(cell, MergedCell):
                    cell.value = answers[no]
        # 判定列に書き込み
        if choices and no in choices:
            for j_col in judgment_cols:
                cell = ws[f"{j_col}{row_idx}"]
                if not isinstance(cell, MergedCell):
                    cell.value = choices[no]

    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))
    wb.close()
    return output_path


def _write_docx_answers(
    original_path: Path,
    answers: dict[int, str],
    config: FormatConfig,
    output_path: Path,
    choices: dict[int, str] | None = None,
    write_cols: list[tuple[str, str]] | None = None,
) -> Path:
    from docx import Document

    doc = Document(str(original_path))
    if config.table_index >= len(doc.tables):
        return output_path

    table = doc.tables[config.table_index]
    q_idx = _col_letter_to_index(config.question_col)

    # 書き込み先の列インデックスを決定
    answer_idxs: list[int] = []
    judgment_idxs: list[int] = []
    if write_cols:
        for col_letter, role in write_cols:
            idx = _col_letter_to_index(col_letter)
            if role == "answer":
                answer_idxs.append(idx)
            elif role == "judgment":
                judgment_idxs.append(idx)
    else:
        answer_idxs = [_col_letter_to_index(config.answer_col)]
        if config.choices_col:
            judgment_idxs = [_col_letter_to_index(config.choices_col)]

    no = 0
    for i, row in enumerate(table.rows):
        if i < config.data_start_row - 1:
            continue
        cells = row.cells
        if len(cells) <= q_idx:
            continue
        q_text = cells[q_idx].text.strip()
        if not q_text:
            continue
        no += 1
        if no in answers:
            for a_idx in answer_idxs:
                if a_idx < len(cells):
                    cells[a_idx].text = answers[no]
        if choices and no in choices:
            for j_idx in judgment_idxs:
                if j_idx < len(cells):
                    cells[j_idx].text = choices[no]

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
